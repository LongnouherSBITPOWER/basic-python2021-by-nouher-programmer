#GUIWiki.py
import wikipedia
#fileที่ได้มาจากไฟล์ python to docx
from docx import Document
def Wiki(keyword,lang='lo'):
    wikipedia.set_lang(lang)
    # summary สำหลับบทความที่สรุป
    data = wikipedia.summary(keyword)
    #page + content บทความทั้งหน้า
    data2 = wikipedia.page(keyword)
    data2 = data2.content
    doc = Document()#สร้างไฟร์ word ใน python
    doc.add_heading(keyword,0)
    doc.add_paragraph(data2)
    doc.save(keyword + '.docx')
    print('สร้างไฟล์สำเร๊จ')

#เปลียนเป็นภาษาไทย
wikipedia.set_lang('th')
from tkinter import *
from tkinter import ttk
from tkinter import messagebox

GUI = Tk()
GUI.title('ບົດຄວາມວິກີພີເດຍ')
GUI.geometry('400x300')
#config
FONT1 = ('Angsana New',15)
#คำอธิบาย
L = ttk.Label(GUI, text='ຄົ້ນຫາບົດຄວາມ',font=FONT1)
L.pack()
#ช่องค้นหาข้อมูล
v_search = StringVar()#กล่องสำหลับเก็บคีเวิร์ด
E1 = ttk.Entry(GUI,textvariable = v_search,font=FONT1,width=40)
E1.pack(pady=10)

#ปุ่มค้นหา

def Search():
    keyword = v_search.get() #.getคือ ดึงข้อมูลเข้ามา
    try:
        #ลองค้นหาดูว่าได้ผลลัพท์หรือไม่ หากได้ให้ผ่านไป
        language = v_radio.get()# th / en / lo
        Wiki(keyword,language)
        messagebox.showinfo('ບັນທຶນຮຽບຮ້ອຍ','ຄົ້ນຫາຂໍ້ຄວາມສຳເລັດບັນທຶນຮຽບຮ້ອຍ')
    except:
        #หากรันคำสั่งเเล้วมีบัญหา เเสดงข้อเเจ้งเตือน
        messagebox.showwarning('keyword Error','กรุนากรอดคำค้นหาไหม่')
        
    #print(wikipedia.search(keyword))
    #result = wikipedia.summary(keyword)
    #print(result)

B1 = ttk.Button(GUI,text='Seach',command=Search)
B1.pack(ipadx=20,ipady=10)#ipadx ขยายในปุ่มเเนวนอน

#เลือกภาษา
F1 = Frame(GUI)
F1.pack(pady=10)

v_radio = StringVar()#=ชองเก็บข้อมูลภาษา

RB1 = ttk.Radiobutton(F1,text='ພາສາລາວ',variable=v_radio,value='lo')
RB2 = ttk.Radiobutton(F1,text='ພາສາອັງກິດ',variable=v_radio,value='en')
RB3 = ttk.Radiobutton(F1,text='ພາສາໄທ',variable=v_radio,value='th')
RB1.invoke()#สั่งให้ค่าเรีมต้นเป็นภาษาไทย

RB1.grid(row=0,column=0)
RB2.grid(row=0,column=1)
RB3.grid(row=0,column=2)





GUI.mainloop()
